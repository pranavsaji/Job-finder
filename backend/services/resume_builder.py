"""
Build a clean, ATS-optimized DOCX resume from structured Claude output.
Single-column layout, standard fonts, no tables/columns/graphics.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name: str, size: float, bold: bool = False, color: tuple = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_section_header(doc: Document, text: str):
    """Add a clean section header with a bottom border line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, "Calibri", 11, bold=True, color=(30, 30, 30))

    # Add bottom border to the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    _set_font(run, "Calibri", 10.5)


def build_ats_resume_docx(data: dict) -> bytes:
    """
    Convert structured resume JSON from Claude into a clean ATS-optimized DOCX.
    Returns bytes of the DOCX file.
    """
    doc = Document()

    # Page margins - standard 1 inch
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Remove default empty paragraph
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    # ---- NAME ----
    name = data.get("candidate_name", "Candidate Name")
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(name)
    _set_font(name_run, "Calibri", 20, bold=True, color=(20, 20, 20))

    # ---- CONTACT LINE ----
    contact = data.get("contact_line", "")
    if contact:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(6)
        c_run = contact_p.add_run(contact)
        _set_font(c_run, "Calibri", 10, color=(80, 80, 80))

    # ---- SUMMARY ----
    summary = data.get("summary", "")
    if summary:
        _add_section_header(doc, "Professional Summary")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        s_run = sp.add_run(summary)
        _set_font(s_run, "Calibri", 10.5)

    # ---- SKILLS ----
    skills = data.get("skills", {})
    tech = skills.get("technical", [])
    tools = skills.get("tools", [])
    soft = skills.get("soft", [])

    if tech or tools or soft:
        _add_section_header(doc, "Skills")
        if tech:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            label = p.add_run("Technical: ")
            _set_font(label, "Calibri", 10.5, bold=True)
            val = p.add_run(", ".join(tech))
            _set_font(val, "Calibri", 10.5)
        if tools:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            label = p.add_run("Tools & Platforms: ")
            _set_font(label, "Calibri", 10.5, bold=True)
            val = p.add_run(", ".join(tools))
            _set_font(val, "Calibri", 10.5)
        if soft:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            label = p.add_run("Core Competencies: ")
            _set_font(label, "Calibri", 10.5, bold=True)
            val = p.add_run(", ".join(soft))
            _set_font(val, "Calibri", 10.5)

    # ---- EXPERIENCE ----
    experience = data.get("experience", [])
    if experience:
        _add_section_header(doc, "Experience")
        for job in experience:
            # Title + Company line
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(1)
            title_run = title_p.add_run(job.get("title", ""))
            _set_font(title_run, "Calibri", 11, bold=True)
            company_str = f"  |  {job.get('company', '')}"
            if job.get("location"):
                company_str += f"  |  {job['location']}"
            co_run = title_p.add_run(company_str)
            _set_font(co_run, "Calibri", 10.5, color=(80, 80, 80))

            # Date line
            start = job.get("start", "")
            end = job.get("end", "Present")
            if start or end:
                date_p = doc.add_paragraph()
                date_p.paragraph_format.space_after = Pt(2)
                date_run = date_p.add_run(f"{start} - {end}")
                _set_font(date_run, "Calibri", 10, color=(100, 100, 100))

            # Bullets
            for bullet in job.get("bullets", []):
                if bullet.strip():
                    _add_bullet(doc, bullet)

    # ---- EDUCATION ----
    education = data.get("education", [])
    if education:
        _add_section_header(doc, "Education")
        for edu in education:
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_before = Pt(4)
            edu_p.paragraph_format.space_after = Pt(1)
            deg_run = edu_p.add_run(edu.get("degree", ""))
            _set_font(deg_run, "Calibri", 11, bold=True)
            school_str = f"  |  {edu.get('school', '')}"
            if edu.get("year"):
                school_str += f"  |  {edu['year']}"
            sc_run = edu_p.add_run(school_str)
            _set_font(sc_run, "Calibri", 10.5, color=(80, 80, 80))

            if edu.get("gpa") or edu.get("honors"):
                detail_p = doc.add_paragraph()
                detail_p.paragraph_format.space_after = Pt(1)
                parts = []
                if edu.get("gpa"):
                    parts.append(f"GPA: {edu['gpa']}")
                if edu.get("honors"):
                    parts.append(str(edu["honors"]))
                d_run = detail_p.add_run(" | ".join(parts))
                _set_font(d_run, "Calibri", 10, color=(100, 100, 100))

    # ---- CERTIFICATIONS ----
    certs = data.get("certifications", [])
    if certs:
        _add_section_header(doc, "Certifications")
        for cert in certs:
            if cert.strip():
                _add_bullet(doc, cert)

    # ---- PROJECTS ----
    projects = data.get("projects", [])
    if projects:
        _add_section_header(doc, "Projects")
        for proj in projects:
            if isinstance(proj, dict) and proj.get("name"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
                name_run = p.add_run(proj["name"])
                _set_font(name_run, "Calibri", 11, bold=True)
                if proj.get("url"):
                    url_run = p.add_run(f"  |  {proj['url']}")
                    _set_font(url_run, "Calibri", 10, color=(60, 120, 200))
                if proj.get("description"):
                    desc_p = doc.add_paragraph()
                    desc_p.paragraph_format.space_after = Pt(1)
                    d_run = desc_p.add_run(proj["description"])
                    _set_font(d_run, "Calibri", 10.5)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
